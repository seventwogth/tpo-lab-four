#!/usr/bin/env python3
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
ASSETS = ROOT / "report_assets"
REPORT = ROOT / "report_lab4.docx"
MAX_RESPONSE_MS = 620
CONFIG_COSTS = {1: 5200, 2: 10100, 3: 10800}


def add_table(doc: Document, df: pd.DataFrame, columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, col in enumerate(columns):
        table.rows[0].cells[idx].text = col
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                cells[idx].text = f"{value:.2f}"
            else:
                cells[idx].text = str(value)


def add_picture_if_exists(doc: Document, filename: str, title: str) -> None:
    path = ASSETS / filename
    if path.exists():
        doc.add_paragraph(title)
        doc.add_picture(str(path), width=Inches(6.4))


def main() -> None:
    load = pd.read_csv(RESULTS / "load_summary.csv")
    stress = pd.read_csv(RESULTS / "stress_summary.csv") if (RESULTS / "stress_summary.csv").exists() else pd.DataFrame()
    selected = None
    selected_path = RESULTS / "selected_config.txt"
    if selected_path.exists():
        selected = int(selected_path.read_text(encoding="utf-8").strip())

    doc = Document()
    doc.add_heading("Лабораторная работа №4", level=0)
    doc.add_paragraph("Нагрузочное и стресс-тестирование веб-приложения с помощью Apache JMeter.")

    doc.add_heading("1. Текст задания", level=1)
    doc.add_paragraph(
        "Вариант: token=497921026, user=-1356108256. Максимальное количество параллельных "
        "пользователей: 7. Средняя нагрузка на одного пользователя: 20 запросов в минуту. "
        "Максимально допустимое время отклика: 620 мс. Требуется проверить конфигурации "
        "$5200, $10100 и $10800, выбрать наиболее дешёвую удовлетворяющую требованиям, "
        "затем провести стресс-тестирование выбранной конфигурации."
    )

    doc.add_heading("2. Описание конфигурации JMeter", level=1)
    doc.add_paragraph("Тестовые планы подготовлены в формате JMX для Apache JMeter 5.6.3 и запускаются через CLI.")

    doc.add_heading("3. Настройки Thread Group", level=1)
    doc.add_paragraph("Нагрузочные тесты: 7 потоков, ramp-up 7 секунд, duration 120 секунд, Loop Forever включён.")

    doc.add_heading("4. Настройки HTTP Request", level=1)
    doc.add_paragraph(
        "Protocol=http, Server=localhost, Port=8079, Method=GET, Path=/. "
        "Параметры запроса: token=497921026, user=-1356108256, config=1/2/3."
    )

    doc.add_heading("5. Assertions", level=1)
    doc.add_paragraph("Response Assertion проверяет HTTP 200. Duration Assertion ограничивает время отклика значением 620 мс.")

    doc.add_heading("6. Constant Throughput Timer", level=1)
    doc.add_paragraph("Для нагрузочного тестирования задан Target Throughput = 140 requests/min, расчёт по всем активным потокам.")

    doc.add_heading("7. Результаты нагрузочного тестирования", level=1)
    load_columns = [
        "config", "cost_usd", "sample_count", "error_count", "error_rate_pct",
        "avg_ms", "median_ms", "p90_ms", "p95_ms", "p99_ms", "max_ms",
        "throughput_rpm", "meets_requirements",
    ]
    add_table(doc, load, load_columns)

    doc.add_heading("8. Графики нагрузочного тестирования", level=1)
    add_picture_if_exists(doc, "load_avg_response_by_config.png", "Среднее время отклика по конфигурациям")
    add_picture_if_exists(doc, "load_response_boxplot.png", "Boxplot распределений времени отклика")
    add_picture_if_exists(doc, "load_percentiles.png", "Перцентили p90/p95/p99")
    add_picture_if_exists(doc, "load_throughput.png", "Throughput")
    add_picture_if_exists(doc, "load_error_rate.png", "Error rate")

    doc.add_heading("9. Выбор конфигурации", level=1)
    if selected is None:
        doc.add_paragraph("По результатам нагрузочного тестирования ни одна конфигурация не удовлетворила требованиям.")
    else:
        row = load[load["config"] == selected].iloc[0]
        doc.add_paragraph(
            f"Выбрана конфигурация {selected} стоимостью ${CONFIG_COSTS[selected]}, так как это наиболее дешёвая "
            f"конфигурация без ошибок и с max response time {row['max_ms']:.2f} мс, что не превышает лимит {MAX_RESPONSE_MS} мс."
        )

    doc.add_heading("10. Настройки стресс-тестирования", level=1)
    doc.add_paragraph(
        "Для выбранной конфигурации выполнены отдельные прогоны: 1/20, 2/40, 3/60, 4/80, "
        "5/100, 7/140, 10/200, 12/240, 16/320 пользователей/requests per minute. "
        "Критерий отказа: max response time > 620 мс или появление HTTP 503."
    )

    doc.add_heading("11. Графики стресс-тестирования", level=1)
    if not stress.empty:
        stress_columns = [
            "config", "users", "target_rpm", "sample_count", "error_count", "http_503_count",
            "error_rate_pct", "avg_ms", "median_ms", "p90_ms", "p95_ms", "p99_ms",
            "max_ms", "throughput_rpm", "failed",
        ]
        add_table(doc, stress, stress_columns)
        add_picture_if_exists(doc, "stress_response_vs_load.png", "Зависимость времени отклика от нагрузки")
        add_picture_if_exists(doc, "stress_error_rate_vs_load.png", "Доля ошибок от нагрузки")
    else:
        doc.add_paragraph("Данные стресс-тестирования отсутствуют.")

    doc.add_heading("12. Анализ результатов", level=1)
    if not stress.empty:
        failed = stress[stress["failed"]].sort_values("target_rpm")
        if failed.empty:
            doc.add_paragraph("В исследованном диапазоне точка отказа не достигнута.")
        else:
            first = failed.iloc[0]
            failed_levels = "; ".join(
                f"{int(row.users)} пользователей / {int(row.target_rpm)} rpm, max={row.max_ms:.2f} мс, HTTP 503={int(row.http_503_count)}"
                for row in failed.itertuples(index=False)
            )
            doc.add_paragraph(
                f"Первое срабатывание критерия отказа зафиксировано на уровне {int(first['users'])} пользователей / "
                f"{int(first['target_rpm'])} requests/min: max response time = {first['max_ms']:.2f} мс, "
                f"HTTP 503 = {int(first['http_503_count'])}."
            )
            doc.add_paragraph(
                "Так как отказ по max response time оказался не монотонным, дополнительно фиксируются все уровни, "
                f"где критерий был нарушен: {failed_levels}."
            )

    doc.add_heading("13. Выводы", level=1)
    if selected is None:
        doc.add_paragraph("Требования варианта не выполнены ни одной из проверенных конфигураций.")
    else:
        doc.add_paragraph(
            f"По результатам реальных прогонов JMeter для дальнейшей эксплуатации выбрана конфигурация {selected}. "
            "Вывод основан на таблицах JTL, сводных CSV и построенных графиках из каталога report_assets."
        )

    doc.save(REPORT)
    print(REPORT.relative_to(ROOT))


if __name__ == "__main__":
    main()
